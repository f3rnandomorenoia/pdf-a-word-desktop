from __future__ import annotations

import io
from collections import Counter
from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips
from pdf2docx import Converter


class ConversionError(RuntimeError):
    """Raised when a PDF conversion cannot be completed."""


def default_output_path(pdf_path: Path, extension: str = "docx") -> Path:
    """Return an output path beside the PDF without overwriting existing files."""
    extension = extension.lower().lstrip(".")
    candidate = pdf_path.with_suffix(f".{extension}")
    if not candidate.exists():
        return candidate

    index = 2
    while True:
        candidate = pdf_path.with_name(f"{pdf_path.stem}-{index}.{extension}")
        if not candidate.exists():
            return candidate
        index += 1


def convert_pdf_to_docx(pdf_path: Path, output_path: Path) -> Path:
    pdf_path = Path(pdf_path)
    output_path = Path(output_path)
    validate_pdf(pdf_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    converter = Converter(str(pdf_path))
    try:
        converter.convert(str(output_path))
    finally:
        converter.close()

    if not output_path.exists() or output_path.stat().st_size == 0:
        raise ConversionError("La conversion no genero un archivo DOCX valido.")
    return output_path


def convert_pdf_to_word(pdf_path: Path, output_path: Path, mode: str = "table") -> Path:
    output_path = Path(output_path)
    extension = output_path.suffix.lower()
    if extension == ".docx":
        if mode == "table":
            return convert_pdf_to_editable_docx(Path(pdf_path), output_path)
        return convert_pdf_to_docx(Path(pdf_path), output_path)
    raise ConversionError("El destino debe terminar en .docx.")


def convert_pdf_to_editable_docx(pdf_path: Path, output_path: Path) -> Path:
    """Build an editable Word document that mirrors the PDF formatting.

    Fonts, sizes, colors, cell shading, border colors, alignments and page
    geometry are extracted from the PDF itself instead of being assumed.
    """
    pdf_path = Path(pdf_path)
    output_path = Path(output_path)
    validate_pdf(pdf_path)

    source = fitz.open(pdf_path)
    try:
        layouts = [_analyze_page(page) for page in source]
        if not any(layout["tables"] or layout["free_lines"] for layout in layouts):
            return convert_pdf_to_docx(pdf_path, output_path)

        content = _content_bbox(layouts, source[0].rect)
        document = Document()
        _configure_document(document, source[0].rect, content, layouts)

        image_cache: dict[int, bytes] = {}
        for index, layout in enumerate(layouts):
            _write_page(document, source, layout, content.y0, first=index == 0, image_cache=image_cache)
    finally:
        source.close()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    if not output_path.exists() or output_path.stat().st_size == 0:
        raise ConversionError("La conversion no genero un archivo DOCX valido.")
    return output_path


# ---------------------------------------------------------------------------
# PDF analysis


def _analyze_page(page: fitz.Page) -> dict:
    tables = []
    for table in page.find_tables().tables:
        if table.row_count == 0 or table.col_count == 0:
            continue
        tables.append({"table": table, "rect": fitz.Rect(table.bbox)})

    spans = _page_spans(page)
    fills: list[tuple[fitz.Rect, tuple]] = []
    strokes: list[tuple[fitz.Rect, tuple, float]] = []
    for drawing in page.get_drawings():
        rect = fitz.Rect(drawing["rect"])
        if drawing["type"] in ("f", "fs") and drawing.get("fill"):
            thickness = min(rect.width, rect.height)
            if thickness <= 1.6:
                strokes.append((rect, drawing["fill"], thickness))
            elif drawing["fill"] != (1.0, 1.0, 1.0):
                fills.append((rect, drawing["fill"]))
        if drawing["type"] in ("s", "fs") and drawing.get("color"):
            strokes.append((rect, drawing["color"], drawing.get("width") or 0.5))

    images: list[dict] = []
    for info in page.get_image_info(xrefs=True):
        xref = info.get("xref")
        if not xref:
            continue
        rect = fitz.Rect(info["bbox"])
        # The PDF may paint several images on the same spot; the last one wins.
        images = [
            image
            for image in images
            if (image["rect"] & rect).get_area() < 0.8 * min(image["rect"].get_area(), rect.get_area())
        ]
        images.append({"rect": rect, "xref": xref})

    table_rects = [entry["rect"] for entry in tables]
    free_spans = [
        span for span in spans if not any(span["rect"].intersects(rect) for rect in table_rects)
    ]
    return {
        "tables": tables,
        "spans": spans,
        "free_lines": _group_lines(free_spans),
        "fills": fills,
        "strokes": strokes,
        "images": images,
    }


def _page_spans(page: fitz.Page) -> list[dict]:
    spans: list[dict] = []
    for block in page.get_text("dict")["blocks"]:
        if block.get("type") != 0:
            continue
        for line in block["lines"]:
            for span in line["spans"]:
                if not span["text"].strip():
                    continue
                spans.append(
                    {
                        "text": span["text"],
                        "font": span["font"],
                        "size": span["size"],
                        "flags": span["flags"],
                        "color": span["color"],
                        "rect": fitz.Rect(span["bbox"]),
                    }
                )

    spans.sort(key=lambda span: (span["rect"].y0, span["rect"].x0))
    return spans


def _group_lines(spans: list[dict]) -> list[dict]:
    """Group spans that share the same baseline into visual lines."""
    spans = sorted(spans, key=lambda span: (span["rect"].y0, span["rect"].x0))
    lines: list[dict] = []
    for span in spans:
        target = None
        for line in reversed(lines[-6:]):
            if _vertical_overlap(line["rect"], span["rect"]) >= 0.5:
                target = line
                break
        if target is None:
            lines.append({"rect": fitz.Rect(span["rect"]), "spans": [span]})
        else:
            target["spans"].append(span)
            target["rect"] |= span["rect"]
    for line in lines:
        line["spans"].sort(key=lambda span: span["rect"].x0)
    lines.sort(key=lambda line: (line["rect"].y0, line["rect"].x0))
    return lines


def _vertical_overlap(a: fitz.Rect, b: fitz.Rect) -> float:
    overlap = min(a.y1, b.y1) - max(a.y0, b.y0)
    smallest = max(min(a.height, b.height), 0.1)
    return overlap / smallest


def _content_bbox(layouts: list[dict], page_rect: fitz.Rect) -> fitz.Rect:
    bbox = fitz.Rect(page_rect.width, page_rect.height, 0, 0)
    found = False
    for layout in layouts:
        for entry in layout["tables"]:
            bbox |= entry["rect"]
            found = True
        for line in layout["free_lines"]:
            bbox |= line["rect"]
            found = True
    if not found:
        return fitz.Rect(36, 36, page_rect.width - 36, page_rect.height - 36)
    return bbox


# ---------------------------------------------------------------------------
# Document construction


def _configure_document(document: Document, page_rect: fitz.Rect, content: fitz.Rect, layouts: list[dict]) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE if page_rect.width > page_rect.height else WD_ORIENT.PORTRAIT
    section.page_width = Pt(page_rect.width)
    section.page_height = Pt(page_rect.height)
    section.left_margin = Pt(max(content.x0, 0))
    section.right_margin = Pt(max(page_rect.width - content.x1, 0))
    section.top_margin = Pt(max(content.y0, 0))
    section.bottom_margin = Pt(12)

    font_counter: Counter = Counter()
    for layout in layouts:
        for span in layout["spans"]:
            family, _, _ = _font_info(span)
            font_counter[(family, _half_points(span["size"]))] += len(span["text"])
    normal_style = document.styles["Normal"]
    if font_counter:
        (family, size), _ = font_counter.most_common(1)[0]
        normal_style.font.name = family
        normal_style.font.size = Pt(size)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)


def _write_page(
    document: Document,
    source: fitz.Document,
    layout: dict,
    top_margin: float,
    first: bool,
    image_cache: dict[int, bytes],
) -> None:
    if not first:
        _add_page_break(document)

    items: list[tuple[str, dict]] = [("table", entry) for entry in layout["tables"]]
    items += [("line", line) for line in layout["free_lines"]]
    items.sort(key=lambda item: item[1]["rect"].y0)

    cursor = top_margin
    previous_kind = None
    for kind, payload in items:
        rect = payload["rect"]
        gap = rect.y0 - cursor
        if kind == "table" and previous_kind == "table":
            _add_spacer(document, max(gap, 1.0))
        elif gap > 0.75:
            _add_spacer(document, gap)

        if kind == "table":
            _add_word_table(document, source, layout, payload, top_margin, image_cache)
        else:
            _add_free_line(document, payload, layout)
        cursor = max(cursor, rect.y1)
        previous_kind = kind


def _add_spacer(document: Document, height_pt: float) -> None:
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    # Word renders an "exact" empty line slightly taller than requested.
    paragraph_format.line_spacing = Pt(max(height_pt - 0.8, 0.75))
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def _add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = Pt(1)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def _add_free_line(document: Document, line: dict, layout: dict) -> None:
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = Pt(max(line["rect"].height, 1))
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    _append_spans(paragraph, line["spans"])


# ---------------------------------------------------------------------------
# Tables


def _add_word_table(
    document: Document,
    source: fitz.Document,
    layout: dict,
    entry: dict,
    left_margin: float,
    image_cache: dict[int, bytes],
) -> None:
    source_table = entry["table"]
    rect = entry["rect"]
    x_edges = _x_edges(source_table)
    y_edges = _y_edges(source_table)
    row_count = len(y_edges) - 1
    col_count = len(x_edges) - 1
    if row_count < 1 or col_count < 1:
        return

    table = document.add_table(rows=row_count, cols=col_count)
    table.style = "Table Grid"
    table.autofit = False
    _set_fixed_table_layout(table)
    border_color, border_size = _border_info(layout, rect)
    _set_table_borders(table, border_color, border_size)
    _set_table_cell_margins(table, top=6, start=16, bottom=6, end=16)

    # Word anchors the table so the first cell text aligns at the margin,
    # shifting it left by the default cell margin; compensate via the indent.
    indent_twips = int(round((rect.x0 - left_margin) * 20)) + 64
    _set_table_indent(table, indent_twips)

    column_widths = [Twips(int(round((x_edges[i + 1] - x_edges[i]) * 20))) for i in range(col_count)]
    # Word renders each "exact" row slightly taller than asked (about half the
    # border width), so compensate to keep the table from drifting down-page.
    border_compensation = border_size * 2.5 / 2 + 0.7
    row_heights = [
        Twips(max(int(round((y_edges[i + 1] - y_edges[i]) * 20 - border_compensation)), 20))
        for i in range(row_count)
    ]
    for col_index, width in enumerate(column_widths):
        table.columns[col_index].width = width
    for row_index, height in enumerate(row_heights):
        table.rows[row_index].height = height
        table.rows[row_index].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    anchors: dict[tuple[int, int], fitz.Rect] = {}
    for row_index, row in enumerate(source_table.rows):
        for col_index, bbox in enumerate(row.cells):
            if bbox is None:
                continue
            start_row, end_row, start_col, end_col = _span_from_bbox(bbox, x_edges, y_edges)
            if (start_row, start_col) in anchors:
                continue
            anchors[(start_row, start_col)] = fitz.Rect(bbox)
            if start_row == end_row and start_col == end_col:
                continue
            anchor = table.cell(start_row, start_col)
            for merge_row in range(start_row, end_row + 1):
                for merge_col in range(start_col, end_col + 1):
                    if merge_row == start_row and merge_col == start_col:
                        continue
                    anchor = anchor.merge(table.cell(merge_row, merge_col))

    for (row_index, col_index), cell_rect in anchors.items():
        cell = table.cell(row_index, col_index)
        start_col = _nearest_index(x_edges, cell_rect.x0)
        end_col = _nearest_index(x_edges, cell_rect.x1) - 1
        span_width = Twips(sum(width.twips for width in column_widths[start_col : end_col + 1]))
        _set_cell_width(cell, span_width)
        _fill_cell(cell, cell_rect, layout, source, image_cache)


def _fill_cell(
    cell,
    cell_rect: fitz.Rect,
    layout: dict,
    source: fitz.Document,
    image_cache: dict[int, bytes],
) -> None:
    fill = _cell_fill(cell_rect, layout["fills"])
    if fill:
        _shade_cell(cell, fill)
    _set_cell_margins(cell, top=6, start=16, bottom=6, end=16)

    cell_spans = [span for span in layout["spans"] if cell_rect.contains(_center(span["rect"]))]
    lines = _group_lines(cell_spans)
    images = [image for image in layout["images"] if cell_rect.contains(_center(image["rect"]))]
    items = [("line", line) for line in lines] + [("image", image) for image in images]
    items.sort(key=lambda item: item[1]["rect"].y0)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    first_paragraph = cell.paragraphs[0]
    _clear_paragraph(first_paragraph)
    for index, (kind, payload) in enumerate(items):
        rect = payload["rect"]
        is_image = kind == "image"
        line_height = rect.height + (2.5 if is_image else 0.4)
        paragraph = first_paragraph if index == 0 else cell.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        if index == 0:
            paragraph_format.space_before = Pt(max(rect.y0 - cell_rect.y0 - 0.7, 0))
        else:
            paragraph_format.space_before = Pt(0)
        if index + 1 < len(items):
            pitch = items[index + 1][1]["rect"].y0 - rect.y0
            paragraph_format.space_after = Pt(max(pitch - line_height, 0))
        else:
            paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = Pt(max(line_height, 1))
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        _align_paragraph(paragraph, rect, cell_rect)

        if kind == "line":
            _append_spans(paragraph, payload["spans"])
        else:
            data = image_cache.get(payload["xref"])
            if data is None:
                data = source.extract_image(payload["xref"])["image"]
                image_cache[payload["xref"]] = data
            run = paragraph.add_run()
            run.add_picture(
                io.BytesIO(data),
                width=Pt(max(rect.width, 1)),
                height=Pt(max(rect.height, 1)),
            )


def _align_paragraph(paragraph, item_rect: fitz.Rect, cell_rect: fitz.Rect, side_margin: float = 0.8) -> None:
    left_gap = item_rect.x0 - cell_rect.x0 - side_margin
    right_gap = cell_rect.x1 - item_rect.x1 - side_margin
    paragraph_format = paragraph.paragraph_format
    if abs(left_gap - right_gap) <= 1.2 and left_gap > 3:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif right_gap < 2 and left_gap > 3:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if right_gap > 1.5:
            paragraph_format.right_indent = Pt(right_gap)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if left_gap > 1.5:
            paragraph_format.left_indent = Pt(left_gap)


def _append_spans(paragraph, spans: list[dict]) -> None:
    previous_x1 = None
    for index, span in enumerate(spans):
        text = span["text"]
        if index == 0:
            text = text.lstrip()
        if index == len(spans) - 1:
            text = text.rstrip()
        if not text:
            continue
        if previous_x1 is not None:
            gap = span["rect"].x0 - previous_x1
            space_width = max(span["size"] * 0.28, 0.5)
            if gap > space_width * 1.5:
                pad_run = paragraph.add_run(" " * min(int(round(gap / space_width)), 80))
                _apply_span_font(pad_run, span)
        run = paragraph.add_run(text)
        _apply_span_font(run, span)
        previous_x1 = span["rect"].x1


def _apply_span_font(run, span: dict) -> None:
    family, bold, italic = _font_info(span)
    run.font.name = family
    run.font.size = Pt(_half_points(span["size"]))
    run.bold = bold
    run.italic = italic
    color = span.get("color") or 0
    if color:
        run.font.color.rgb = RGBColor((color >> 16) & 255, (color >> 8) & 255, color & 255)


def _font_info(span: dict) -> tuple[str, bool, bool]:
    raw = span["font"].split("+")[-1]
    lowered = raw.lower()
    bold = bool(span["flags"] & 16) or "bold" in lowered
    italic = bool(span["flags"] & 2) or "italic" in lowered or "oblique" in lowered
    family = raw.split("-")[0].split(",")[0]
    if family.endswith("MT"):
        family = family[:-2]
    if family.endswith("PS"):
        family = family[:-2]
    return family or "Arial", bold, italic


def _half_points(size: float) -> float:
    return max(round(size * 2) / 2, 1)


def _cell_fill(cell_rect: fitz.Rect, fills: list[tuple[fitz.Rect, tuple]]) -> str | None:
    cell_area = max(cell_rect.get_area(), 0.1)
    best: tuple[float, tuple] | None = None
    for rect, color in fills:
        overlap = fitz.Rect(rect) & cell_rect
        ratio = overlap.get_area() / cell_area
        if ratio >= 0.5 and (best is None or ratio > best[0]):
            best = (ratio, color)
    if best is None:
        return None
    return _color_hex(best[1])


def _border_info(layout: dict, table_rect: fitz.Rect) -> tuple[str, int]:
    pad = fitz.Rect(table_rect.x0 - 2, table_rect.y0 - 2, table_rect.x1 + 2, table_rect.y1 + 2)
    counter: Counter = Counter()
    widths: list[float] = []
    for rect, color, width in layout["strokes"]:
        if pad.intersects(rect):
            counter[color] += 1
            widths.append(width)
    if not counter:
        return "000000", 4
    color = counter.most_common(1)[0][0]
    widths.sort()
    median = widths[len(widths) // 2] if widths else 0.5
    size = min(max(int(round(median * 8)), 2), 24)
    return _color_hex(color), size


def _color_hex(color: tuple) -> str:
    return "".join(f"{int(round(channel * 255)):02X}" for channel in color[:3])


def _center(rect: fitz.Rect) -> fitz.Point:
    return fitz.Point((rect.x0 + rect.x1) / 2, (rect.y0 + rect.y1) / 2)


def _nearest_index(edges: list[float], value: float) -> int:
    return min(range(len(edges)), key=lambda index: abs(edges[index] - value))


def _span_from_bbox(
    bbox: tuple[float, float, float, float], x_edges: list[float], y_edges: list[float]
) -> tuple[int, int, int, int]:
    x0, y0, x1, y1 = bbox
    start_col = _nearest_index(x_edges, x0)
    end_col = max(_nearest_index(x_edges, x1) - 1, start_col)
    start_row = _nearest_index(y_edges, y0)
    end_row = max(_nearest_index(y_edges, y1) - 1, start_row)
    return start_row, end_row, start_col, end_col


def _x_edges(source_table) -> list[float]:
    return sorted({round(cell[0], 3) for cell in source_table.cells} | {round(cell[2], 3) for cell in source_table.cells})


def _y_edges(source_table) -> list[float]:
    return sorted({round(cell[1], 3) for cell in source_table.cells} | {round(cell[3], 3) for cell in source_table.cells})


def _clear_paragraph(paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        if child.tag == qn("w:pPr"):
            continue
        element.remove(child)


# ---------------------------------------------------------------------------
# Low level DOCX helpers


def _set_fixed_table_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_table_cell_margins(table, top: int, start: int, bottom: int, end: int) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.first_child_found_in("w:tblCellMar")
    if existing is not None:
        tbl_pr.remove(existing)
    margins = OxmlElement("w:tblCellMar")
    for tag, value in (("top", top), ("left", start), ("bottom", bottom), ("right", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)


def _set_table_indent(table, twips: int) -> None:
    tbl_pr = table._tbl.tblPr
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(twips))
    indent.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str, size: int) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.first_child_found_in("w:tblBorders")
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    tbl_pr.append(borders)


def _set_cell_width(cell, width) -> None:
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width.twips))


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def validate_pdf(pdf_path: Path) -> None:
    if not pdf_path.exists():
        raise ConversionError("El archivo PDF no existe.")
    if not pdf_path.is_file():
        raise ConversionError("La ruta seleccionada no es un archivo.")
    if pdf_path.suffix.lower() != ".pdf":
        raise ConversionError("Selecciona un archivo con extension .pdf.")
    if pdf_path.stat().st_size == 0:
        raise ConversionError("El PDF esta vacio.")
